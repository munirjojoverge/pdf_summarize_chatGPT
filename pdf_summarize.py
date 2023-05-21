# ********************************************************
#                      Munir Jojo-Verge (c)
#                         May 21th 2023
# ********************************************************

import os
import glob
import pdfplumber
from docx import Document
import openai
import hydra
from hydra.utils import get_original_cwd
from omegaconf import DictConfig
import logging
from tqdm import tqdm
import tiktoken

# Configure logging
log = logging.getLogger(__name__) 


def count_tokens(text):
    """
    Count the number of tokens in the given text.

    Args:
        text (str): Text to count tokens in.

    Returns:
        int: Number of tokens in the text.
    """
    enc = tiktoken.encoding_for_model("text-davinci-002")
    tokens = enc.encode_ordinary(text)
    return len(tokens)


@hydra.main(config_path="config", config_name="config")
def main(cfg: DictConfig) -> None:
    """
    Summarize PDFs using GPT-3.

    Args:
        cfg (DictConfig): Hydra configuration object.
    """
    cwd = get_original_cwd()
    pdfs_path = os.path.join(cwd, cfg.file.pdfs_path)
    output_path = os.path.join(cwd, cfg.file.output_path)

    # Set the API key
    openai.api_key = cfg.openai.api_key    
    engine = cfg.openai.engine    
    temperature = cfg.openai.temperature
    max_tokens = cfg.openai.max_tokens
    stop = cfg.openai.stop
    token_limit = cfg.openai.token_limit
    buffer_tokens = cfg.openai.buffer_tokens

    # Loop over all PDFs
    for filename in tqdm(glob.glob(os.path.join(pdfs_path, '*.pdf')), desc="Processing PDFs"):
        # Create a new Word document for each PDF
        doc = Document()

        with pdfplumber.open(filename) as pdf:            
            # Initialize variables
            summary = ""
            tokens_read = 0
            pages_processed = 0

            # Iterate over the pages_text until reaching the token limit
            with tqdm(total=len(pdf.pages), desc="Processing Pages") as pbar_pages:
                while tokens_read < token_limit and pages_processed < len(pdf.pages):
                    # Extract text from the current page
                    page_text = pdf.pages[pages_processed].extract_text()

                    # Calculate the number of tokens in the page
                    page_tokens = count_tokens(page_text)

                    # Check if adding the page's tokens exceeds the token limit
                    if tokens_read + page_tokens + buffer_tokens > token_limit:
                        break

                    # Add the page's text to the summary
                    summary += page_text.strip() + "\n"

                    # Update the tokens read count
                    tokens_read += page_tokens

                    # Increment the page counter
                    pages_processed += 1

                    # Update the progress bar for pages
                    pbar_pages.update(1)

            # Split the summary into smaller segments to fit within the context length limit
            segment_size = token_limit - buffer_tokens
            segments = [summary[i:i+segment_size] for i in range(0, len(summary), segment_size)]

            # Initialize the final summary
            final_summary = ""

            # Process each segment and generate summaries
            for segment in segments:
                # Ask GPT to summarize the segment
                response = openai.Completion.create(
                    engine=engine,
                    prompt=f"{segment}\nSummarize this section of a paper for further analysis and code development. " + 
                    "Include the key findings, methodology, experimental results, statistical analysis, and any limitations. " + 
                    "Ensure the summary is clear, concise, and coherent. Provide context about the field of study and research question. " +
                    "Format the summary with bullet points for easy reference.",
                    temperature=temperature,
                    max_tokens=max_tokens,
                    stop=stop,
                )

                # Add the summary to the final summary
                final_summary += response.choices[0].text.strip() + "\n"

            # Create a new Word document for each PDF
            doc = Document()

            # Add the final summary to the document
            doc.add_paragraph(final_summary)

            # Save the document with the same name as the PDF file
            output_filename = os.path.splitext(os.path.basename(filename))[0] + '_summary.docx'
            output_filename = os.path.join(output_path, output_filename)
            doc.save(output_filename)

if __name__ == "__main__":
    main()
